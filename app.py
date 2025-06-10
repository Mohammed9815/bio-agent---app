import streamlit as st
import pandas as pd
from io import BytesIO
from zipfile import ZipFile
import base64
import random

# --- استيراد المكتبات ---
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import arabic_reshaper
from bidi.algorithm import get_display

# --- إعدادات الصفحة ---
st.set_page_config(page_title="الوكيل الذكي لمادة الأحياء", layout="wide", page_icon="🧬")

# --- العناوين الرئيسية ---
st.markdown('<h1 style="text-align:center;">🧬 الوكيل الذكي 5.0 🧬</h1>', unsafe_allow_html=True)
st.markdown('<h4 style="text-align:center;">مساعدك الشخصي لتوليد أنشطة طلابية فريدة ومبتكرة</h4>', unsafe_allow_html=True)
st.markdown("<hr/>", unsafe_allow_html=True)

# ==============================================================================
#  المرحلة الثانية: إعادة هيكلة الدروس بشكل كامل
# ==============================================================================
LESSONS_DB = {
    "الفصل الدراسي الأول": {
        "الوحدة الأولى: تركيب الخلية": [
            "١-١ علم الخلية واستخدام المجهر",
            "١-٢ الخلايا النباتية والخلايا الحيوانية كما تُرى بالمجهر الضوئي",
            "١-٣ حساب القياسات ومقدار التكبير",
            "١-٤ الخلايا النباتية والخلايا الحيوانية كما تُرى بالمجهر الإلكتروني",
            "١-٥ البكتيريا",
            "١-٦ الفيروسات"
        ],
        "الوحدة الثانية: الجزيئات الحيوية": [
            "٢-١ الكيمياء الحيوية",
            "٢-٢ الكربوهيدرات",
            "٢-٣ الدهون",
            "٢-٤ البروتينات",
            "٢-٥ الماء"
        ],
        "الوحدة الثالثة: الإنزيمات": [
            "٣-١ ما هو الإنزيم؟",
            "٣-٢ طريقة عمل الإنزيمات",
            "٣-٣ استقصاء سير تفاعل محفز بالإنزيم",
            "٣-٤ العوامل المؤثرة في عمل الإنزيم",
            "٣-٥ مقارنة ألفة (تلاؤم) الإنزيمات",
            "٣-٦ مثبطات الإنزيم",
            "٣-٧ الإنزيمات المثبتة"
        ],
        "الوحدة الرابعة: دورة الخلية والانقسام المتساوي": [
            "٤-١ النمو والتكاثر و دور الكروموسومات",
            "٤-٢ دورة الخلية",
            "٤-٣ الانقسام المتساوي",
            "٤-٤ دور التيلوميرات",
            "٤-٥ دور الخلايا الجذعية",
            "٤-٦ السرطانات"
        ]
    },
    "الفصل الدراسي الثاني": {
        "الوحدة الخامسة: أغشية الخلية والنقل": [
            "٥-١ وظائف الأغشية وتركيبها",
            "٥-٢ وظائف الجزيئات الموجودة في الأغشية",
            "٥-٣ التأشير الخلوي",
            "٥-٤ حركة المواد عبر الأغشية"
        ],
        "الوحدة السادسة: النقل في النباتات": [
            "٦-١ تركيب السيقان والجذور والأوراق وتوزيع نسيجي الخشب واللحاء",
            "٦-٢ نقل الماء",
            "٦-٣ نقل نواتج التمثيل الغذائي"
        ],
        "الوحدة السابعة: النقل في الثدييات": [
            "٧-١ الأوعية الدموية",
            "٧-٢ السائل النسيجي",
            "٧-٣ الدم",
            "٧-٤ القلب"
        ],
        "الوحدة الثامنة: تبادل الغازات": [
            "٨-١ الرئتان",
            "٨-٢ تدفئة وتنظيف الهواء",
            "٨-٣ الحويصلات الهوائية"
        ]
    }
}


# --- بنك الأنشطة (لا تغيير) ---
ACTIVITY_BANK = {
    "علاجي": [
        "اكتب تعريفاً مبسطاً لمفهوم '{lesson}'.",
        "صل بين المصطلحات التالية وما يناسبها من تعاريف بخصوص درس '{lesson}'. (سيقوم المعلم بتوفير المصطلحات)",
        "أكمل الفراغ: من أهم أجزاء '{lesson}' هي ______ و ______. (مثال توضيحي)",
        "ارسم شكلاً مبسطاً يوضح فكرة '{lesson}' مع كتابة البيانات الأساسية.",
        "اذكر وظيفة واحدة رئيسية لـ '{lesson}' في جسم الكائن الحي."
    ],
    "دعم": [
        "لخص في ثلاث نقاط أهم الأفكار في درس '{lesson}'.",
        "قارن بين مفهومين مرتبطين بدرس '{lesson}' (مثلاً: الانتشار البسيط والانتشار المسهل).",
        "اشرح لزميل لك كيف تعمل الآلية الخاصة بـ '{lesson}'.",
        "حلل الرسم البياني أو الشكل الموجود في الكتاب المدرسي صفحة (X) المتعلق بدرس '{lesson}'.",
        "صمم خريطة مفاهيمية بسيطة توضح العلاقات بين المكونات الرئيسية لـ '{lesson}'."
    ],
    "إثرائي": [
        "ابحث عن مرض أو حالة طبية ترتبط بخلل في آلية '{lesson}' واكتب فقرة موجزة عنها.",
        "اقترح طريقة مبتكرة لشرح مفهوم '{lesson}' باستخدام مواد بسيطة من الحياة اليومية.",
        "ماذا سيحدث لو لم تكن عملية '{lesson}' موجودة؟ صف التأثيرات المحتملة.",
        "ابحث عن أحدث الاكتشافات العلمية المتعلقة بـ '{lesson}' خلال السنوات الخمس الماضية.",
        "صمم سؤالاً واحداً بمستوى تفكير عليا (تحليل، تركيب، تقويم) حول درس '{lesson}' مع نموذج إجابته."
    ]
}

# --- وظائف مساعدة (لا تغيير) ---
def generate_smart_activity(score):
    if score < 5:
        level = "علاجي"
        level_emoji = "😕"
    elif 5 <= score <= 7:
        level = "دعم"
        level_emoji = "💪"
    else:
        level = "إثرائي"
        level_emoji = "😃"
    activity_template = random.choice(ACTIVITY_BANK[level])
    return f"{level} {level_emoji}", activity_template

def create_word_doc(name, level, content):
    document = Document()
    for section in document.sections:
        section.right_to_left = True
    def add_rtl_paragraph(text, alignment=WD_ALIGN_PARAGRAPH.RIGHT, size=12, bold=False):
        reshaped_text = arabic_reshaper.reshape(text)
        bidi_text = get_display(reshaped_text)
        p = document.add_paragraph()
        p.alignment = alignment
        run = p.add_run(bidi_text)
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(size)
        font.bold = bold
        p_format = p.paragraph_format
        p_format.right_to_left = True
    add_rtl_paragraph("الوكيل الذكي لمادة الأحياء", alignment=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    add_rtl_paragraph(f"اسم الطالب: {name}", size=14)
    add_rtl_paragraph(f"التصنيف: {level}", size=14)
    document.add_paragraph("--------------------------------------------------")
    for line in content.split('\n'):
        add_rtl_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# --- واجهة المستخدم المطورة ---
df = None
with st.container(border=True):
    st.subheader("📥 الخطوة 1: أدخل بيانات الطلاب")
    method = st.radio("اختر طريقة الإدخال:", ["📄 رفع ملف Excel", "✍️ إدخال يدوي"], horizontal=True, label_visibility="collapsed")

    if method == "📄 رفع ملف Excel":
        excel_file = st.file_uploader("ارفع ملف Excel (يحتوي على عمودي 'الاسم' و 'الدرجة')", type=["xlsx"])
        if excel_file:
            df = pd.read_excel(excel_file)
    else:
        count = st.number_input("حدد عدد الطلاب:", min_value=1, max_value=50, value=1, step=1)
        data = {'الاسم': [], 'الدرجة': []}
        for i in range(count):
            c1, c2 = st.columns([3, 1])
            with c1:
                name = st.text_input(f"اسم الطالب {i+1}", key=f"n{i}")
            with c2:
                score = st.number_input("الدرجة", 0.0, 10.0, 7.0, step=0.1, key=f"s{i}")
            data['الاسم'].append(name)
            data['الدرجة'].append(score)
        df = pd.DataFrame(data)
        df = df[df['الاسم'].str.strip() != ""] if not df.empty else df

with st.container(border=True):
    st.subheader("📚 الخطوة 2: اختر الدرس")
    
    # الواجهة الجديدة لاختيار الدروس
    selected_semester = st.selectbox("اختر الفصل الدراسي:", list(LESSONS_DB.keys()))
    
    if selected_semester:
        units = list(LESSONS_DB[selected_semester].keys())
        selected_unit = st.selectbox(f"اختر الوحدة الدراسية لـ {selected_semester}:", units)
        
        if selected_unit:
            lessons = LESSONS_DB[selected_semester][selected_unit]
            selected_lesson = st.selectbox(f"اختر الدرس من {selected_unit}:", lessons)

# --- المنطق الرئيسي ---
if df is not None and not df.empty and 'الاسم' in df.columns and 'الدرجة' in df.columns and 'selected_lesson' in locals():
    if st.button("✨ توليد الأنشطة الذكية", use_container_width=True, type="primary"):
        with st.spinner('الوكيل الذكي يفكر... 🧠 لطفاً، انتظر قليلاً.'):
            files_to_zip = []
            st.markdown("<hr/>", unsafe_allow_html=True)
            st.subheader("📋 النتائج والأنشطة المخصصة")

            for index, row in df.iterrows():
                name, score = row['الاسم'], row['الدرجة']
                if pd.notna(name) and name.strip() != "" and pd.notna(score):
                    level, activity_template = generate_smart_activity(float(score))
                    final_activity = activity_template.format(lesson=selected_lesson)
                    
                    with st.expander(f"👤 {name}  |  الدرجة: {score}  |  المستوى المقترح: {level}"):
                        st.text_area("النشاط المولد:", final_activity, height=150, key=f"activity_{index}")
                        word_buffer = create_word_doc(name, level, final_activity)
                        files_to_zip.append((f"{name}.docx", word_buffer.getvalue()))

            if files_to_zip:
                zip_buf = BytesIO()
                with ZipFile(zip_buf, "w") as zipf:
                    for filename, data in files_to_zip:
                        zipf.writestr(filename, data)
                zip_buf.seek(0)
                b64 = base64.b64encode(zip_buf.read()).decode()
                download_filename = f"أنشطة_{selected_lesson.replace(' ', '_')}.zip"
                
                st.markdown("<br>", unsafe_allow_html=True)
                st.download_button(
                    label="📥 تحميل الأنشطة",
                    data=zip_buf,
                    file_name=download_filename,
                    mime="application/zip",
                    use_container_width=True
                )
        st.success("🎉 تم توليد الأنشطة بنجاح!")
        st.balloons()

